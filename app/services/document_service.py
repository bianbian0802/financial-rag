"""Service layer responsible for validating and storing uploaded documents."""

import json
import logging
import re
from datetime import UTC, datetime
from pathlib import Path
from urllib.parse import urlparse
from uuid import uuid4

import httpx
from docx import Document as WordDocument
from fastapi import UploadFile
from pydantic import ValidationError
from pypdf import PdfReader

from app.core.config import Settings
from app.core.exceptions import AppException
from app.schemas.document import (
    DocumentChunk,
    DocumentChunkResponse,
    DocumentEmbeddingResponse,
    EmbeddedDocumentChunk,
    EmbeddingUsage,
    DocumentParseResponse,
    DocumentUploadResponse,
    ParsedDocumentRecord,
    StoredDocumentMetadata,
)

logger = logging.getLogger(__name__)


class DocumentService:
    """Validate supported uploads and persist them to local storage in chunks."""

    supported_extensions = {".pdf", ".md", ".txt", ".doc", ".docx"}
    metadata_directory_name = "metadata"
    parsed_directory_name = "parsed"
    chunks_directory_name = "chunks"
    embeddings_directory_name = "embeddings"
    preview_text_limit = 240

    def __init__(self, settings: Settings) -> None:
        """Store document-related settings used by the upload workflow."""
        self.settings = settings

    async def save_upload(self, upload_file: UploadFile) -> DocumentUploadResponse:
        """Persist a supported uploaded file to disk using bounded memory."""
        self._validate_filename(upload_file.filename)

        original_filename = Path(upload_file.filename or "").name
        file_extension = Path(original_filename).suffix.lower()
        document_id = uuid4().hex
        stored_filename = f"{document_id}{file_extension}"
        storage_directory = Path(self.settings.documents_storage_dir)
        storage_directory.mkdir(parents=True, exist_ok=True)
        storage_path = storage_directory / stored_filename

        logger.info("Saving uploaded document %s as %s.", original_filename, stored_filename)

        max_size_bytes = self.settings.documents_max_upload_size_mb * 1024 * 1024
        chunk_size_bytes = self.settings.documents_chunk_size_bytes
        written_size_bytes = 0

        try:
            with storage_path.open("wb") as output_file:
                while True:
                    chunk = await upload_file.read(chunk_size_bytes)
                    if not chunk:
                        break

                    written_size_bytes += len(chunk)
                    if written_size_bytes > max_size_bytes:
                        raise AppException(
                            message=(
                                f"Uploaded file is too large. Maximum allowed size is "
                                f"{self.settings.documents_max_upload_size_mb} MB."
                            ),
                            status_code=413,
                            error_code="DOCUMENT_TOO_LARGE",
                            details={"max_size_mb": self.settings.documents_max_upload_size_mb},
                        )

                    output_file.write(chunk)
        except AppException:
            self._remove_partial_file(storage_path)
            raise
        except OSError as exc:
            self._remove_partial_file(storage_path)
            raise AppException(
                message="Failed to store the uploaded document.",
                status_code=500,
                error_code="DOCUMENT_STORE_FAILED",
            ) from exc
        finally:
            await upload_file.close()

        logger.info(
            "Document upload completed for %s with %s bytes written.",
            original_filename,
            written_size_bytes,
        )

        upload_response = DocumentUploadResponse(
            document_id=document_id,
            original_filename=original_filename,
            stored_filename=stored_filename,
            file_extension=file_extension,
            content_type=upload_file.content_type or "application/octet-stream",
            size_bytes=written_size_bytes,
            storage_path=str(storage_path.as_posix()),
            status="uploaded",
            uploaded_at=datetime.now(UTC),
        )
        self._write_document_metadata(upload_response)
        return upload_response

    def parse_document(self, document_id: str) -> DocumentParseResponse:
        """Load an uploaded document by ID, extract plain text, and persist the parse result."""
        metadata = self._load_document_metadata(document_id)
        storage_path = Path(metadata.storage_path)
        if not storage_path.exists():
            raise AppException(
                message="The uploaded document file could not be found on disk.",
                status_code=404,
                error_code="DOCUMENT_FILE_MISSING",
                details={"document_id": document_id},
            )

        parser_name, extracted_text = self._extract_text(storage_path, metadata.file_extension)
        normalized_text = self._normalize_extracted_text(extracted_text)
        if not normalized_text:
            raise AppException(
                message="The document was parsed, but no extractable text was found.",
                status_code=422,
                error_code="DOCUMENT_TEXT_EMPTY",
                details={"document_id": document_id, "file_extension": metadata.file_extension},
            )

        parsed_directory = self._get_parsed_directory()
        parsed_directory.mkdir(parents=True, exist_ok=True)
        parsed_output_path = parsed_directory / f"{document_id}.json"
        parsed_at = datetime.now(UTC)

        parsed_record = ParsedDocumentRecord(
            document_id=metadata.document_id,
            original_filename=metadata.original_filename,
            stored_filename=metadata.stored_filename,
            file_extension=metadata.file_extension,
            source_storage_path=metadata.storage_path,
            parser_name=parser_name,
            status="parsed",
            extracted_char_count=len(normalized_text),
            preview_text=self._build_preview_text(normalized_text),
            parsed_output_path=str(parsed_output_path.as_posix()),
            parsed_at=parsed_at,
            extracted_text=normalized_text,
        )
        self._write_json_file(parsed_output_path, parsed_record.model_dump(mode="json"))

        logger.info(
            "Document %s parsed successfully using %s with %s extracted characters.",
            document_id,
            parser_name,
            parsed_record.extracted_char_count,
        )

        return DocumentParseResponse.model_validate(parsed_record.model_dump())

    def chunk_document(self, document_id: str) -> DocumentChunkResponse:
        """Load a parsed document, clean the text, and split it into retrieval chunks."""
        parsed_record = self._load_parsed_document(document_id)
        cleaned_text = self._clean_text_for_chunking(parsed_record.extracted_text)
        if not cleaned_text:
            raise AppException(
                message="The parsed document does not contain usable text for chunking.",
                status_code=422,
                error_code="DOCUMENT_CHUNK_EMPTY",
                details={"document_id": document_id},
            )

        chunk_size = self.settings.document_chunk_size
        chunk_overlap = self.settings.document_chunk_overlap
        self._validate_chunk_settings(chunk_size, chunk_overlap)

        chunks = self._split_text_into_chunks(
            document_id=document_id,
            cleaned_text=cleaned_text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        chunks_directory = self._get_chunks_directory()
        chunks_directory.mkdir(parents=True, exist_ok=True)
        chunks_output_path = chunks_directory / f"{document_id}.json"
        chunked_at = datetime.now(UTC)

        chunk_response = DocumentChunkResponse(
            document_id=document_id,
            source_parsed_output_path=parsed_record.parsed_output_path,
            chunk_count=len(chunks),
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            cleaned_char_count=len(cleaned_text),
            chunks_output_path=str(chunks_output_path.as_posix()),
            status="chunked",
            chunked_at=chunked_at,
            chunks=chunks,
        )
        self._write_json_file(chunks_output_path, chunk_response.model_dump(mode="json"))

        logger.info(
            "Document %s chunked successfully into %s chunks with size=%s overlap=%s.",
            document_id,
            len(chunks),
            chunk_size,
            chunk_overlap,
        )

        return chunk_response

    async def embed_document(self, document_id: str) -> DocumentEmbeddingResponse:
        """Load chunked text for a document, generate embeddings, and persist the vectors."""
        chunk_record = self._load_chunked_document(document_id)
        if not chunk_record.chunks:
            raise AppException(
                message="The chunked document does not contain any chunks to embed.",
                status_code=422,
                error_code="DOCUMENT_EMBED_EMPTY",
                details={"document_id": document_id},
            )

        self._validate_embedding_settings()

        embedded_chunks: list[EmbeddedDocumentChunk] = []
        total_prompt_tokens = 0
        total_tokens = 0
        batch_size = self.settings.embedding_batch_size
        chunk_group_size = max(1, batch_size)

        logger.info(
            "Embedding document %s with %s chunks using batch size %s.",
            document_id,
            len(chunk_record.chunks),
            chunk_group_size,
        )

        for batch_start in range(0, len(chunk_record.chunks), chunk_group_size):
            batch = chunk_record.chunks[batch_start : batch_start + chunk_group_size]
            embeddings, usage = await self._request_embeddings([chunk.text for chunk in batch])
            if len(embeddings) != len(batch):
                raise AppException(
                    message="Embedding provider returned an unexpected number of vectors.",
                    status_code=502,
                    error_code="DOCUMENT_EMBEDDING_RESPONSE_INVALID",
                    details={
                        "expected_count": len(batch),
                        "received_count": len(embeddings),
                    },
                )

            if usage is not None:
                total_prompt_tokens += usage.prompt_tokens
                total_tokens += usage.total_tokens

            for chunk, embedding in zip(batch, embeddings, strict=True):
                embedded_chunks.append(
                    EmbeddedDocumentChunk(
                        chunk_id=chunk.chunk_id,
                        chunk_index=chunk.chunk_index,
                        text=chunk.text,
                        char_count=chunk.char_count,
                        start_char_index=chunk.start_char_index,
                        end_char_index=chunk.end_char_index,
                        preview_text=chunk.preview_text,
                        embedding=embedding,
                        embedding_dimensions=len(embedding),
                    )
                )

        embedding_dimensions = embedded_chunks[0].embedding_dimensions
        embeddings_directory = self._get_embeddings_directory()
        embeddings_directory.mkdir(parents=True, exist_ok=True)
        embeddings_output_path = embeddings_directory / f"{document_id}.json"
        embedded_at = datetime.now(UTC)
        usage_payload = None
        if total_prompt_tokens or total_tokens:
            usage_payload = EmbeddingUsage(
                prompt_tokens=total_prompt_tokens,
                total_tokens=total_tokens,
            )

        embedding_response = DocumentEmbeddingResponse(
            document_id=document_id,
            source_chunks_output_path=chunk_record.chunks_output_path,
            embedding_model=self._get_embedding_model(),
            provider="openai-compatible",
            chunk_count=len(embedded_chunks),
            embedding_dimensions=embedding_dimensions,
            embeddings_output_path=str(embeddings_output_path.as_posix()),
            status="embedded",
            usage=usage_payload,
            embedded_at=embedded_at,
            embedded_chunks=embedded_chunks,
        )
        self._write_json_file(embeddings_output_path, embedding_response.model_dump(mode="json"))

        logger.info(
            "Document %s embedded successfully into %s vectors with dimension %s.",
            document_id,
            len(embedded_chunks),
            embedding_dimensions,
        )

        return embedding_response

    def _validate_filename(self, filename: str | None) -> None:
        """Ensure the client provided a supported filename before reading any content."""
        if not filename:
            raise AppException(
                message="Uploaded file must include a filename.",
                status_code=400,
                error_code="DOCUMENT_FILENAME_MISSING",
            )

        normalized_extension = Path(filename).suffix.lower()
        if normalized_extension not in self.supported_extensions:
            raise AppException(
                message=(
                    "Unsupported document type. Supported extensions are: "
                    ".pdf, .md, .txt, .doc, .docx."
                ),
                status_code=415,
                error_code="DOCUMENT_TYPE_UNSUPPORTED",
                details={"supported_extensions": sorted(self.supported_extensions)},
            )

    def _write_document_metadata(self, metadata: StoredDocumentMetadata) -> None:
        """Persist upload metadata so later steps can resolve documents by document ID."""
        metadata_directory = self._get_metadata_directory()
        metadata_directory.mkdir(parents=True, exist_ok=True)
        metadata_path = metadata_directory / f"{metadata.document_id}.json"
        self._write_json_file(metadata_path, metadata.model_dump(mode="json"))

    def _load_document_metadata(self, document_id: str) -> StoredDocumentMetadata:
        """Load persisted upload metadata, reconstructing it for older uploads when needed."""
        if not document_id:
            raise AppException(
                message="Document ID must not be empty.",
                status_code=400,
                error_code="DOCUMENT_ID_MISSING",
            )

        metadata_path = self._get_metadata_directory() / f"{document_id}.json"
        if metadata_path.exists():
            try:
                metadata_payload = json.loads(metadata_path.read_text(encoding="utf-8"))
                return StoredDocumentMetadata.model_validate(metadata_payload)
            except (OSError, json.JSONDecodeError, ValidationError) as exc:
                raise AppException(
                    message="Stored document metadata is corrupted or unreadable.",
                    status_code=500,
                    error_code="DOCUMENT_METADATA_INVALID",
                    details={"document_id": document_id},
                ) from exc

        return self._reconstruct_metadata(document_id)

    def _reconstruct_metadata(self, document_id: str) -> StoredDocumentMetadata:
        """Rebuild minimal metadata for files uploaded before metadata persistence existed."""
        storage_directory = Path(self.settings.documents_storage_dir)
        matching_files = [
            candidate
            for candidate in storage_directory.glob(f"{document_id}.*")
            if candidate.is_file() and candidate.suffix.lower() in self.supported_extensions
        ]
        if not matching_files:
            raise AppException(
                message="No uploaded document was found for the provided document ID.",
                status_code=404,
                error_code="DOCUMENT_NOT_FOUND",
                details={"document_id": document_id},
            )

        source_file = matching_files[0]
        source_file_stat = source_file.stat()
        reconstructed_metadata = StoredDocumentMetadata(
            document_id=document_id,
            original_filename=source_file.name,
            stored_filename=source_file.name,
            file_extension=source_file.suffix.lower(),
            content_type="application/octet-stream",
            size_bytes=source_file_stat.st_size,
            storage_path=str(source_file.as_posix()),
            status="uploaded",
            uploaded_at=datetime.fromtimestamp(source_file_stat.st_mtime, tz=UTC),
        )
        logger.warning(
            "Reconstructed document metadata for %s because no metadata file existed.",
            document_id,
        )
        self._write_document_metadata(reconstructed_metadata)
        return reconstructed_metadata

    def _load_parsed_document(self, document_id: str) -> ParsedDocumentRecord:
        """Load the persisted parsed document record produced by the Day10 parse step."""
        parsed_output_path = self._get_parsed_directory() / f"{document_id}.json"
        if not parsed_output_path.exists():
            raise AppException(
                message="No parsed document was found for the provided document ID.",
                status_code=404,
                error_code="DOCUMENT_PARSED_NOT_FOUND",
                details={"document_id": document_id},
            )

        try:
            parsed_payload = json.loads(parsed_output_path.read_text(encoding="utf-8"))
            return ParsedDocumentRecord.model_validate(parsed_payload)
        except (OSError, json.JSONDecodeError, ValidationError) as exc:
            raise AppException(
                message="Stored parsed document output is corrupted or unreadable.",
                status_code=500,
                error_code="DOCUMENT_PARSED_INVALID",
                details={"document_id": document_id},
            ) from exc

    def _load_chunked_document(self, document_id: str) -> DocumentChunkResponse:
        """Load the persisted chunk document record produced by the Day11 chunk step."""
        chunks_output_path = self._get_chunks_directory() / f"{document_id}.json"
        if not chunks_output_path.exists():
            raise AppException(
                message="No chunked document was found for the provided document ID.",
                status_code=404,
                error_code="DOCUMENT_CHUNKED_NOT_FOUND",
                details={"document_id": document_id},
            )

        try:
            chunk_payload = json.loads(chunks_output_path.read_text(encoding="utf-8"))
            return DocumentChunkResponse.model_validate(chunk_payload)
        except (OSError, json.JSONDecodeError, ValidationError) as exc:
            raise AppException(
                message="Stored chunked document output is corrupted or unreadable.",
                status_code=500,
                error_code="DOCUMENT_CHUNKED_INVALID",
                details={"document_id": document_id},
            ) from exc

    def _extract_text(self, storage_path: Path, file_extension: str) -> tuple[str, str]:
        """Dispatch to the appropriate parser based on the stored file extension."""
        if file_extension in {".md", ".txt"}:
            return "plain_text", self._read_text_document(storage_path)
        if file_extension == ".pdf":
            return "pypdf", self._extract_pdf_text(storage_path)
        if file_extension == ".docx":
            return "python-docx", self._extract_docx_text(storage_path)
        if file_extension == ".doc":
            raise AppException(
                message="Legacy .doc parsing is not supported yet. Please convert the file to .docx.",
                status_code=422,
                error_code="DOCUMENT_PARSE_UNSUPPORTED",
                details={"file_extension": ".doc"},
            )

        raise AppException(
            message="Unsupported document type for parsing.",
            status_code=415,
            error_code="DOCUMENT_PARSE_TYPE_UNSUPPORTED",
            details={"file_extension": file_extension},
        )

    @staticmethod
    def _read_text_document(storage_path: Path) -> str:
        """Read plain-text documents using a small set of practical encodings."""
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return storage_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
            except OSError as exc:
                raise AppException(
                    message="Failed to read the uploaded text document.",
                    status_code=500,
                    error_code="DOCUMENT_READ_FAILED",
                ) from exc

        raise AppException(
            message="Failed to decode the uploaded text document with the supported encodings.",
            status_code=422,
            error_code="DOCUMENT_DECODE_FAILED",
            details={"supported_encodings": ["utf-8", "utf-8-sig", "gb18030"]},
        )

    @staticmethod
    def _extract_pdf_text(storage_path: Path) -> str:
        """Extract text from each PDF page and merge the results into one string."""
        try:
            reader = PdfReader(str(storage_path))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise AppException(
                message="Failed to extract text from the uploaded PDF document.",
                status_code=422,
                error_code="DOCUMENT_PARSE_FAILED",
                details={"file_extension": ".pdf"},
            ) from exc

    @staticmethod
    def _extract_docx_text(storage_path: Path) -> str:
        """Extract paragraph and table cell text from a DOCX document."""
        try:
            document = WordDocument(str(storage_path))
        except Exception as exc:
            raise AppException(
                message="Failed to open the uploaded DOCX document.",
                status_code=422,
                error_code="DOCUMENT_PARSE_FAILED",
                details={"file_extension": ".docx"},
            ) from exc

        fragments = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    fragments.append(" | ".join(row_cells))
        return "\n".join(fragments)

    @staticmethod
    def _normalize_extracted_text(extracted_text: str) -> str:
        """Normalize line endings and collapse noisy blank lines in parsed output."""
        normalized_text = extracted_text.replace("\r\n", "\n").replace("\r", "\n")
        normalized_text = "\n".join(line.rstrip() for line in normalized_text.split("\n"))
        normalized_text = re.sub(r"\n{3,}", "\n\n", normalized_text)
        return normalized_text.strip()

    @staticmethod
    def _clean_text_for_chunking(extracted_text: str) -> str:
        """Apply lightweight cleanup so chunk boundaries are more stable and useful."""
        cleaned_text = extracted_text.replace("\r\n", "\n").replace("\r", "\n")
        cleaned_text = re.sub(r"[ \t]+", " ", cleaned_text)
        cleaned_text = re.sub(r"\n[ \t]+", "\n", cleaned_text)
        cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)
        return cleaned_text.strip()

    @staticmethod
    def _validate_chunk_settings(chunk_size: int, chunk_overlap: int) -> None:
        """Ensure chunking settings are valid before generating offsets and overlaps."""
        if chunk_size <= 0:
            raise AppException(
                message="Document chunk size must be greater than zero.",
                status_code=500,
                error_code="DOCUMENT_CHUNK_CONFIG_INVALID",
                details={"document_chunk_size": chunk_size},
            )
        if chunk_overlap < 0 or chunk_overlap >= chunk_size:
            raise AppException(
                message="Document chunk overlap must be zero or greater and smaller than chunk size.",
                status_code=500,
                error_code="DOCUMENT_CHUNK_CONFIG_INVALID",
                details={
                    "document_chunk_size": chunk_size,
                    "document_chunk_overlap": chunk_overlap,
                },
            )

    def _split_text_into_chunks(
        self,
        *,
        document_id: str,
        cleaned_text: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> list[DocumentChunk]:
        """Split cleaned text into overlapping chunks, preferring paragraph-like boundaries."""
        chunks: list[DocumentChunk] = []
        start_index = 0
        text_length = len(cleaned_text)
        step_size = chunk_size - chunk_overlap

        while start_index < text_length:
            tentative_end = min(start_index + chunk_size, text_length)
            end_index = self._find_chunk_boundary(cleaned_text, start_index, tentative_end)
            chunk_text = cleaned_text[start_index:end_index].strip()
 
            if not chunk_text:
                start_index = min(start_index + step_size, text_length)
                continue

            chunk_index = len(chunks)
            chunks.append(
                DocumentChunk(
                    chunk_id=f"{document_id}-chunk-{chunk_index:04d}",
                    chunk_index=chunk_index,
                    text=chunk_text,
                    char_count=len(chunk_text),
                    start_char_index=start_index,
                    end_char_index=end_index,
                    preview_text=self._build_chunk_preview_text(chunk_text),
                )
            )

            if end_index >= text_length:
                break

            next_start_index = max(end_index - chunk_overlap, start_index + 1)
            start_index = next_start_index

        return chunks

    @staticmethod
    def _find_chunk_boundary(cleaned_text: str, start_index: int, tentative_end: int) -> int:
        """Prefer paragraph and sentence boundaries so chunks read more naturally."""
        if tentative_end >= len(cleaned_text):
            return len(cleaned_text)

        boundary_window = cleaned_text[start_index:tentative_end]
        for marker in ("\n\n", "\n", "。", "！", "？", ".", "!", "?", "；", ";", "，", ",", " "):
            boundary_index = boundary_window.rfind(marker)
            if boundary_index > 0:
                return start_index + boundary_index + len(marker)

        return tentative_end

    def _build_chunk_preview_text(self, chunk_text: str) -> str:
        """Build a short preview string for chunk-level response payloads."""
        preview_limit = self.settings.document_chunk_preview_limit
        if len(chunk_text) <= preview_limit:
            return chunk_text
        return f"{chunk_text[:preview_limit].rstrip()}..."

    def _validate_embedding_settings(self) -> None:
        """Ensure embedding configuration is present before calling the provider."""
        if not self._get_embedding_base_url():
            raise AppException(
                message="Embedding base URL is missing. Please configure embedding_base_url or llm_base_url.",
                status_code=500,
                error_code="EMBEDDING_BASE_URL_MISSING",
            )
        if not self._get_embedding_model():
            raise AppException(
                message="Embedding model is missing. Please configure embedding_model in your environment.",
                status_code=500,
                error_code="EMBEDDING_MODEL_MISSING",
            )
        if not self._get_embedding_api_key() and not self._is_local_base_url(self._get_embedding_base_url()):
            raise AppException(
                message="Embedding API key is missing. Please configure embedding_api_key in your environment.",
                status_code=500,
                error_code="EMBEDDING_API_KEY_MISSING",
            )
        if self.settings.embedding_batch_size <= 0:
            raise AppException(
                message="Embedding batch size must be greater than zero.",
                status_code=500,
                error_code="EMBEDDING_BATCH_SIZE_INVALID",
                details={"embedding_batch_size": self.settings.embedding_batch_size},
            )

    async def _request_embeddings(self, texts: list[str]) -> tuple[list[list[float]], EmbeddingUsage | None]:
        """Call the configured OpenAI-compatible embeddings endpoint for a batch of texts."""
        try:
            async with httpx.AsyncClient(
                timeout=self.settings.embedding_timeout_seconds,
                headers=self._build_embedding_headers(),
            ) as client:
                response = await client.post(
                    self._build_embedding_url(),
                    json={
                        "model": self._get_embedding_model(),
                        "input": texts,
                    },
                )
                response.raise_for_status()
        except httpx.TimeoutException as exc:
            raise AppException(
                message="Embedding request timed out.",
                status_code=504,
                error_code="EMBEDDING_TIMEOUT",
            ) from exc
        except httpx.HTTPStatusError as exc:
            logger.warning("Embedding provider returned error: %s", exc.response.text)
            raise AppException(
                message="Embedding provider returned an error.",
                status_code=502,
                error_code="EMBEDDING_PROVIDER_ERROR",
                details={"status_code": exc.response.status_code},
            ) from exc
        except httpx.HTTPError as exc:
            raise AppException(
                message="Failed to connect to the embedding provider.",
                status_code=502,
                error_code="EMBEDDING_CONNECTION_ERROR",
            ) from exc

        try:
            response_data = response.json()
        except ValueError as exc:
            raise AppException(
                message="Embedding provider returned a non-JSON response.",
                status_code=502,
                error_code="EMBEDDING_RESPONSE_INVALID",
            ) from exc
        return self._extract_embedding_vectors(response_data)

    @staticmethod
    def _extract_embedding_vectors(response_data: dict) -> tuple[list[list[float]], EmbeddingUsage | None]:
        """Extract embedding vectors and optional usage metadata from provider output."""
        try:
            data = response_data["data"]
        except (KeyError, TypeError) as exc:
            raise AppException(
                message="Embedding response format is invalid.",
                status_code=502,
                error_code="EMBEDDING_RESPONSE_INVALID",
            ) from exc

        embeddings: list[list[float]] = []
        try:
            for item in sorted(data, key=lambda entry: entry["index"]):
                embedding = item["embedding"]
                embeddings.append([float(value) for value in embedding])
        except (KeyError, TypeError, ValueError) as exc:
            raise AppException(
                message="Embedding response format is invalid.",
                status_code=502,
                error_code="EMBEDDING_RESPONSE_INVALID",
            ) from exc

        usage = None
        raw_usage = response_data.get("usage")
        if isinstance(raw_usage, dict) and "prompt_tokens" in raw_usage and "total_tokens" in raw_usage:
            try:
                usage = EmbeddingUsage(
                    prompt_tokens=int(raw_usage["prompt_tokens"]),
                    total_tokens=int(raw_usage["total_tokens"]),
                )
            except (TypeError, ValueError) as exc:
                raise AppException(
                    message="Embedding response usage format is invalid.",
                    status_code=502,
                    error_code="EMBEDDING_USAGE_INVALID",
                ) from exc

        return embeddings, usage

    def _build_embedding_headers(self) -> dict[str, str]:
        """Build HTTP headers for the embedding provider request."""
        headers = {"Content-Type": "application/json"}
        api_key = self._get_embedding_api_key()
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        return headers

    def _build_embedding_url(self) -> str:
        """Return the full embedding endpoint URL."""
        return f"{self._get_embedding_base_url().rstrip('/')}/embeddings"

    def _get_embedding_base_url(self) -> str:
        """Resolve the embedding provider base URL with chat URL as a fallback."""
        return self.settings.embedding_base_url or self.settings.llm_base_url

    def _get_embedding_api_key(self) -> str:
        """Resolve the embedding provider API key with chat key as a fallback."""
        return self.settings.embedding_api_key or self.settings.llm_api_key

    def _get_embedding_model(self) -> str:
        """Return the configured embedding model."""
        return self.settings.embedding_model

    @staticmethod
    def _is_local_base_url(base_url: str) -> bool:
        """Check whether the configured provider points to a local development host."""
        hostname = urlparse(base_url).hostname
        return hostname in {"127.0.0.1", "localhost"}

    def _build_preview_text(self, extracted_text: str) -> str:
        """Create a short preview for API responses without returning the full parsed text."""
        if len(extracted_text) <= self.preview_text_limit:
            return extracted_text
        return f"{extracted_text[: self.preview_text_limit].rstrip()}..."

    def _write_json_file(self, output_path: Path, payload: dict[str, object]) -> None:
        """Write structured JSON to disk using UTF-8 so later pipeline steps can reuse it."""
        try:
            output_path.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except OSError as exc:
            raise AppException(
                message="Failed to persist document metadata or parsed output.",
                status_code=500,
                error_code="DOCUMENT_PERSIST_FAILED",
                details={"output_path": str(output_path)},
            ) from exc

    def _get_metadata_directory(self) -> Path:
        """Return the directory used for upload metadata sidecar files."""
        return Path(self.settings.documents_storage_dir) / self.metadata_directory_name

    def _get_parsed_directory(self) -> Path:
        """Return the directory used for parsed text output files."""
        return Path(self.settings.documents_storage_dir) / self.parsed_directory_name

    def _get_chunks_directory(self) -> Path:
        """Return the directory used for parsed document chunk output files."""
        return Path(self.settings.documents_storage_dir) / self.chunks_directory_name

    def _get_embeddings_directory(self) -> Path:
        """Return the directory used for parsed document embedding output files."""
        return Path(self.settings.documents_storage_dir) / self.embeddings_directory_name

    @staticmethod
    def _remove_partial_file(storage_path: Path) -> None:
        """Delete a partially written file if the upload fails midway."""
        if storage_path.exists():
            storage_path.unlink()
