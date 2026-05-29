"""Tests for the Day10 document parsing API."""

import io
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document as WordDocument
from fastapi.testclient import TestClient

from app.api.v1.documents import get_document_service
from app.core.config import Settings
from app.main import app
from app.services.document_service import DocumentService


class DocumentParsingApiTests(unittest.TestCase):
    """Verify uploaded documents can be parsed into plain text through the API."""

    def setUp(self) -> None:
        """Create isolated document storage and a fresh API client for each test."""
        self.temp_directory = tempfile.TemporaryDirectory()
        self.client = TestClient(app)
        self._override_document_service()

    def tearDown(self) -> None:
        """Remove dependency overrides and temporary files after each test."""
        app.dependency_overrides.clear()
        self.temp_directory.cleanup()

    def test_parse_markdown_document_succeeds(self) -> None:
        """Ensure a stored markdown document can be parsed into plain text."""
        upload_response = self.client.post(
            "/api/v1/documents/upload",
            files={"file": ("guide.md", b"# Financial RAG\nDay10 parsing works.", "text/markdown")},
        )
        document_id = upload_response.json()["document_id"]

        parse_response = self.client.post(f"/api/v1/documents/{document_id}/parse")

        self.assertEqual(parse_response.status_code, 200)
        payload = parse_response.json()
        self.assertEqual(payload["document_id"], document_id)
        self.assertEqual(payload["parser_name"], "plain_text")
        self.assertEqual(payload["status"], "parsed")
        self.assertIn("Day10 parsing works.", payload["preview_text"])
        parsed_output_path = Path(payload["parsed_output_path"])
        self.assertTrue(parsed_output_path.exists())
        stored_payload = json.loads(parsed_output_path.read_text(encoding="utf-8"))
        self.assertIn("Financial RAG", stored_payload["extracted_text"])

    def test_parse_docx_document_succeeds(self) -> None:
        """Ensure DOCX uploads can be parsed with paragraph content intact."""
        upload_response = self.client.post(
            "/api/v1/documents/upload",
            files={
                "file": (
                    "report.docx",
                    self._build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        document_id = upload_response.json()["document_id"]

        parse_response = self.client.post(f"/api/v1/documents/{document_id}/parse")

        self.assertEqual(parse_response.status_code, 200)
        payload = parse_response.json()
        self.assertEqual(payload["parser_name"], "python-docx")
        self.assertIn("Revenue increased by 12 percent.", payload["preview_text"])
        stored_payload = json.loads(Path(payload["parsed_output_path"]).read_text(encoding="utf-8"))
        self.assertIn("Quarter", stored_payload["extracted_text"])

    def test_parse_pdf_document_succeeds(self) -> None:
        """Ensure PDF uploads can be parsed with extractable page text."""
        upload_response = self.client.post(
            "/api/v1/documents/upload",
            files={"file": ("statement.pdf", self._build_pdf_bytes("Hello PDF parser"), "application/pdf")},
        )
        document_id = upload_response.json()["document_id"]

        parse_response = self.client.post(f"/api/v1/documents/{document_id}/parse")

        self.assertEqual(parse_response.status_code, 200)
        payload = parse_response.json()
        self.assertEqual(payload["parser_name"], "pypdf")
        self.assertIn("Hello PDF parser", payload["preview_text"])

    def test_parse_legacy_doc_returns_clear_error(self) -> None:
        """Ensure legacy .doc files return a clear unsupported parsing error."""
        upload_response = self.client.post(
            "/api/v1/documents/upload",
            files={"file": ("legacy.doc", b"fake binary doc", "application/msword")},
        )
        document_id = upload_response.json()["document_id"]

        parse_response = self.client.post(f"/api/v1/documents/{document_id}/parse")

        self.assertEqual(parse_response.status_code, 422)
        payload = parse_response.json()
        self.assertEqual(payload["error_code"], "DOCUMENT_PARSE_UNSUPPORTED")

    def test_parse_missing_document_returns_not_found(self) -> None:
        """Ensure parsing an unknown document ID returns a stable not-found response."""
        parse_response = self.client.post("/api/v1/documents/not-real/parse")

        self.assertEqual(parse_response.status_code, 404)
        payload = parse_response.json()
        self.assertEqual(payload["error_code"], "DOCUMENT_NOT_FOUND")

    def _override_document_service(self) -> None:
        """Override the document service dependency with isolated test storage settings."""
        settings = Settings(
            documents_storage_dir=self.temp_directory.name,
            documents_max_upload_size_mb=25,
            documents_chunk_size_bytes=1024 * 1024,
        )
        app.dependency_overrides[get_document_service] = lambda: DocumentService(settings)

    @staticmethod
    def _build_docx_bytes() -> bytes:
        """Create a small DOCX document in memory for integration-style parsing tests."""
        buffer = io.BytesIO()
        document = WordDocument()
        document.add_heading("Quarterly Financial Summary", level=1)
        document.add_paragraph("Revenue increased by 12 percent.")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Net profit"
        table.rows[1].cells[1].text = "8.6M"
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _build_pdf_bytes(text: str) -> bytes:
        """Create a tiny PDF with extractable text without introducing extra test dependencies."""
        escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT\n/F1 18 Tf\n72 720 Td\n({escaped_text}) Tj\nET"
        objects = [
            "<< /Type /Catalog /Pages 2 0 R >>",
            "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            (
                "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
            ),
            f"<< /Length {len(stream.encode('latin-1'))} >>\nstream\n{stream}\nendstream",
            "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]

        pdf_parts = ["%PDF-1.4\n"]
        object_offsets: list[int] = []
        current_offset = len(pdf_parts[0].encode("latin-1"))
        for index, object_body in enumerate(objects, start=1):
            object_text = f"{index} 0 obj\n{object_body}\nendobj\n"
            object_offsets.append(current_offset)
            pdf_parts.append(object_text)
            current_offset += len(object_text.encode("latin-1"))

        xref_offset = current_offset
        pdf_parts.append(f"xref\n0 {len(objects) + 1}\n")
        pdf_parts.append("0000000000 65535 f \n")
        for object_offset in object_offsets:
            pdf_parts.append(f"{object_offset:010d} 00000 n \n")
        pdf_parts.append(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n")
        pdf_parts.append(f"startxref\n{xref_offset}\n%%EOF")
        return "".join(pdf_parts).encode("latin-1")


if __name__ == "__main__":
    unittest.main()
